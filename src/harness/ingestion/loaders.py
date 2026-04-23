"""Document loaders for different file types."""

from __future__ import annotations

import logging
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Supported file-type mappings
_EXT_MAP: dict[str, str] = {
    ".pdf": "pdf",
    ".html": "html",
    ".htm": "html",
    ".md": "markdown",
    ".markdown": "markdown",
    ".docx": "docx",
    ".doc": "docx",
    ".csv": "csv",
    ".tsv": "csv",
    ".txt": "text",
    ".text": "text",
    ".rst": "text",
    ".log": "text",
}


@dataclass
class Document:
    """A loaded document with text content and metadata.

    Attributes:
        id:          Unique identifier (UUID hex).
        content:     Full extracted text.
        metadata:    Source-specific metadata (title, page_count, etc.).
        source_path: Original path or URL.
        file_type:   Detected file type (pdf, html, markdown, docx, csv, text).
    """

    id: str
    content: str
    metadata: dict
    source_path: str
    file_type: str = "text"

    @classmethod
    def create(
        cls,
        content: str,
        source_path: str,
        file_type: str,
        metadata: Optional[dict] = None,
    ) -> "Document":
        """Factory method that auto-generates an ID."""
        return cls(
            id=uuid.uuid4().hex,
            content=content,
            metadata=metadata or {},
            source_path=source_path,
            file_type=file_type,
        )


# ---------------------------------------------------------------------------
# Per-type loaders
# ---------------------------------------------------------------------------


async def load_pdf(path: str) -> list[Document]:
    """Load a PDF file, extracting text per page using PyMuPDF (fitz).

    Skips encrypted PDFs and logs extraction errors per page.

    Args:
        path: Filesystem path to the PDF file.

    Returns:
        List of Document objects, one per non-empty page.
    """
    try:
        import fitz  # type: ignore  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF (fitz) not installed; cannot load PDF: %s", path)
        return []

    docs: list[Document] = []
    try:
        pdf = fitz.open(path)
    except Exception as exc:
        logger.error("Failed to open PDF %s: %s", path, exc)
        return []

    if pdf.is_encrypted:
        logger.warning("Skipping encrypted PDF: %s", path)
        pdf.close()
        return []

    page_count = len(pdf)
    for page_num in range(page_count):
        try:
            page = pdf[page_num]
            text = page.get_text("text")
            if not text.strip():
                continue
            docs.append(
                Document.create(
                    content=text,
                    source_path=path,
                    file_type="pdf",
                    metadata={
                        "source": path,
                        "page": page_num + 1,
                        "page_count": page_count,
                        "title": pdf.metadata.get("title", "") if pdf.metadata else "",
                    },
                )
            )
        except Exception as exc:
            logger.warning("Error extracting page %d from %s: %s", page_num + 1, path, exc)

    pdf.close()
    return docs


async def load_html(url_or_path: str) -> list[Document]:
    """Load an HTML file or URL, extracting main text content via trafilatura.

    Args:
        url_or_path: A file path or HTTP/HTTPS URL.

    Returns:
        List containing one Document with the extracted main content.
    """
    try:
        import trafilatura  # type: ignore
    except ImportError:
        logger.warning("trafilatura not installed; falling back to plain HTML strip")
        trafilatura = None  # type: ignore

    html_content: Optional[str] = None

    # Check if it's a URL
    if url_or_path.startswith(("http://", "https://")):
        try:
            import httpx  # type: ignore
            async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
                response = await client.get(url_or_path)
                response.raise_for_status()
                html_content = response.text
        except Exception as exc:
            logger.error("Failed to fetch URL %s: %s", url_or_path, exc)
            return []
    else:
        p = Path(url_or_path)
        if not p.exists():
            logger.error("HTML file not found: %s", url_or_path)
            return []
        html_content = p.read_text(encoding="utf-8", errors="replace")

    if not html_content:
        return []

    # Extract main content
    if trafilatura is not None:
        extracted = trafilatura.extract(
            html_content,
            include_comments=False,
            include_tables=True,
            no_fallback=False,
        )
        text = extracted or html_content
    else:
        # Minimal HTML stripping fallback
        import re
        text = re.sub(r"<[^>]+>", " ", html_content)
        text = re.sub(r"\s+", " ", text).strip()

    if not text.strip():
        return []

    return [
        Document.create(
            content=text,
            source_path=url_or_path,
            file_type="html",
            metadata={"source": url_or_path},
        )
    ]


async def load_markdown(path: str) -> list[Document]:
    """Load a Markdown file, splitting into sections by ## headers.

    Args:
        path: Filesystem path to the .md file.

    Returns:
        List of Documents, one per top-level section (or one if no headers).
    """
    import re

    p = Path(path)
    if not p.exists():
        logger.error("Markdown file not found: %s", path)
        return []

    content = p.read_text(encoding="utf-8", errors="replace")
    if not content.strip():
        return []

    # Split on ## (H2) headers; keep the content together with the header
    parts = re.split(r"(?m)^(#{1,2} .+)$", content)

    docs: list[Document] = []
    current_title = p.stem
    current_text: list[str] = []

    for part in parts:
        part_stripped = part.strip()
        if re.match(r"^#{1,2} ", part_stripped):
            # Save previous section if non-empty
            if current_text:
                body = "\n".join(current_text).strip()
                if body:
                    docs.append(
                        Document.create(
                            content=body,
                            source_path=path,
                            file_type="markdown",
                            metadata={"source": path, "section": current_title},
                        )
                    )
            current_title = part_stripped.lstrip("#").strip()
            current_text = [part_stripped]
        else:
            current_text.append(part_stripped)

    # Last section
    if current_text:
        body = "\n".join(current_text).strip()
        if body:
            docs.append(
                Document.create(
                    content=body,
                    source_path=path,
                    file_type="markdown",
                    metadata={"source": path, "section": current_title},
                )
            )

    if not docs:
        # No headers: return as single document
        docs.append(
            Document.create(
                content=content,
                source_path=path,
                file_type="markdown",
                metadata={"source": path},
            )
        )

    return docs


async def load_docx(path: str) -> list[Document]:
    """Load a .docx file, extracting paragraphs and table text.

    Args:
        path: Filesystem path to the .docx file.

    Returns:
        List containing one Document with all paragraphs and table contents.
    """
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError:
        logger.warning("python-docx not installed; cannot load .docx: %s", path)
        return []

    p = Path(path)
    if not p.exists():
        logger.error("DOCX file not found: %s", path)
        return []

    try:
        doc = DocxDocument(str(p))
    except Exception as exc:
        logger.error("Failed to open DOCX %s: %s", path, exc)
        return []

    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table cells
    for table_idx, table in enumerate(doc.tables):
        parts.append(f"[Table {table_idx + 1}]")
        for row in table.rows:
            row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_parts:
                parts.append(" | ".join(row_parts))

    content = "\n".join(parts)
    if not content.strip():
        return []

    return [
        Document.create(
            content=content,
            source_path=path,
            file_type="docx",
            metadata={
                "source": path,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )
    ]


async def load_csv(path: str) -> list[Document]:
    """Load a CSV file and convert each row to a text description.

    Each row is converted to: "Row N: col1=val1, col2=val2, ..."

    Args:
        path: Filesystem path to the .csv file.

    Returns:
        List containing one Document with all rows as text.
    """
    import csv

    p = Path(path)
    if not p.exists():
        logger.error("CSV file not found: %s", path)
        return []

    rows: list[str] = []
    try:
        with p.open("r", encoding="utf-8-sig", errors="replace", newline="") as fh:
            # Auto-detect delimiter
            sample = fh.read(4096)
            fh.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",\t|;")
            except csv.Error:
                dialect = csv.excel  # type: ignore
            reader = csv.DictReader(fh, dialect=dialect)
            headers = reader.fieldnames or []
            for row_num, row in enumerate(reader, 1):
                parts = [f"{k}={v}" for k, v in row.items() if v is not None]
                rows.append(f"Row {row_num}: " + ", ".join(parts))
    except Exception as exc:
        logger.error("Failed to read CSV %s: %s", path, exc)
        return []

    if not rows:
        return []

    content = "\n".join(rows)
    return [
        Document.create(
            content=content,
            source_path=path,
            file_type="csv",
            metadata={"source": path, "row_count": len(rows), "columns": headers},
        )
    ]


async def load_text(path: str) -> list[Document]:
    """Load a plain text file.

    Args:
        path: Filesystem path to the text file.

    Returns:
        List containing one Document with the full file content.
    """
    p = Path(path)
    if not p.exists():
        logger.error("Text file not found: %s", path)
        return []

    try:
        content = p.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        logger.error("Failed to read text file %s: %s", path, exc)
        return []

    if not content.strip():
        return []

    return [
        Document.create(
            content=content,
            source_path=path,
            file_type="text",
            metadata={"source": path, "size_bytes": p.stat().st_size},
        )
    ]


# ---------------------------------------------------------------------------
# Auto-detect and dispatch
# ---------------------------------------------------------------------------


def detect_type(path: str) -> str:
    """Return the file type string for a given path based on its extension.

    Args:
        path: File path (or URL with extension).

    Returns:
        One of: "pdf", "html", "markdown", "docx", "csv", "text".
        Defaults to "text" for unknown extensions.
    """
    suffix = Path(path).suffix.lower()
    return _EXT_MAP.get(suffix, "text")


async def load(path: str) -> list[Document]:
    """Auto-detect file type and dispatch to the appropriate loader.

    Args:
        path: File path or URL to load.

    Returns:
        List of Document objects.
    """
    # Check for URL patterns regardless of extension
    if path.startswith(("http://", "https://")):
        return await load_html(path)

    file_type = detect_type(path)
    loader_map = {
        "pdf": load_pdf,
        "html": load_html,
        "markdown": load_markdown,
        "docx": load_docx,
        "csv": load_csv,
        "text": load_text,
    }
    loader = loader_map.get(file_type, load_text)
    return await loader(path)
